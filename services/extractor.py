import io
import re
from typing import List, Optional
import pdfplumber
import docx

def clean_text(text: str) -> str:
    """Normalizes whitespace, removes control characters, and formats bullet points."""
    if not text:
        return ""
    # Normalize unicode non-breaking spaces
    text = text.replace('\xa0', ' ')
    # Normalize bullet markers
    text = re.sub(r'[\u2022\u2023\u25e6\u2043\u2219]', '- ', text)
    # Remove excessive blank lines (3+ to 2)
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Strip trailing whitespace per line
    lines = [line.strip() for line in text.splitlines()]
    return '\n'.join(lines).strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text from PDF bytes using pdfplumber."""
    extracted_text = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text(layout=True)
            if page_text:
                extracted_text.append(page_text)
            else:
                # Fallback to standard extraction
                fallback_text = page.extract_text()
                if fallback_text:
                    extracted_text.append(fallback_text)
                    
    raw_content = "\n\n".join(extracted_text)
    return clean_text(raw_content)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX bytes including paragraphs and table contents."""
    doc = docx.Document(io.BytesIO(file_bytes))
    extracted_text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            extracted_text.append(para.text.strip())
            
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                extracted_text.append(" | ".join(row_cells))
                
    raw_content = "\n".join(extracted_text)
    return clean_text(raw_content)


def extract_text_from_file(filename: str, file_bytes: bytes) -> str:
    """Dispatches text extraction based on file extension."""
    ext = filename.lower().split('.')[-1]
    if ext == 'pdf':
        return extract_text_from_pdf(file_bytes)
    elif ext in ['docx', 'doc']:
        return extract_text_from_docx(file_bytes)
    elif ext in ['txt', 'md']:
        return clean_text(file_bytes.decode('utf-8', errors='ignore'))
    else:
        raise ValueError(f"Unsupported file format: .{ext}. Supported formats: .pdf, .docx, .txt")
